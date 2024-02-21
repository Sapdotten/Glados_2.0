import docx
from docx.shared import Pt

doc = docx.Document("C:\\Users\\levtr\\OneDrive\\Рабочий стол\\Glados_2.0\\modules\\user\\SampleThanks.docx")
#стили шрифта
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(14)

#Обучающимся института
institute = doc.paragraphs[3]
run = institute.add_run(input("Введите институт: ") + ":")
run.bold = True

#Заполнение таблицы
FIO = doc.tables[2]

while True:
    num_rows = len(FIO.rows)
    cells = FIO.add_row().cells
    cells[0].text = str(num_rows) + (".") #порядковый номер
    cells[1].text = input("Введите ФИО: ") #ФИО
    cells[2].text = input("Введите Группу: ") #Группа
    exit = input("Чтобы закончить напишите 'E': ")
    if exit.lower() == "e":
        break

#Организация и проведение
for_what = doc.paragraphs[6]
run = for_what.add_run(input("за организацию и проведение: ") + ".")

doc.save("C:\\Users\\levtr\\OneDrive\\Рабочий стол\\Glados_2.0\\modules\\user\\ReadyFile\\ReadySampleThanks.docx")
