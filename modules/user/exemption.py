from docx import Document
from docx.shared import Pt


def exemption():
    doc = Document("SampleExemption.docx")  # макет документа в папке с кодом

    # стили шрифта
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # добавление института в строку ворд
    institute = doc.paragraphs[1]
    institute.add_run(input("Введите институт: "))

    # Ввод И. О. Фамилии
    iof = doc.paragraphs[2]
    iof.add_run(input("Введите И.О.Фамилию: "))

    # Ввод названия мероприятия
    event_name = input("Название мероприятия: ")

    for i in doc.paragraphs:
        i.text = i.text.replace("event", event_name)

    # Ввод даты мероприятия
    date_event = input("Дата мероприятия: ")

    for j in doc.paragraphs:
        j.text = j.text.replace("date", date_event)

    # Заполнение таблицы
    fio = doc.tables[0]
    while True:
        num_rows = len(fio.rows)
        cells = fio.add_row().cells
        cells[0].text = str(num_rows) + (".")  # порядковый номер
        cells[1].text = input("Введите ФИО: ")  # ФИО
        cells[2].text = input("Введите Группу: ")  # Группа
        next_step = input("Для продолжения нажмите Enter, Чтобы закончить напишите 'e', : ")
        if next_step.lower() == "e" or next_step.lower() == "е":  # e на русском и на англ
            break

    # Сохранение документа
    doc.save(
        f"ReadyExemption\\Освобождение {event_name} {date_event}.docx")  # необходимо создать папку "ReadyExemption" в корне проекта


exemption()
