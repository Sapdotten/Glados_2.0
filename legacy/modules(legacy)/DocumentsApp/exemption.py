from docx import Document
from docx.shared import Pt
from utils.config_manager import DocumentsConfigs


class Documents:
    exemption_doc: Document
    thanks_doc: Document
    _sample_dir: str
    _doc_style = 'Normal'
    _font_style = 'Times New Roman'
    _font_size = 14

    def __init__(self):
        self.date = 'date'
        self.event_name = 'event name'
        self.institut_name = 'institut name'
        self.director_name = 'director name'
        self.fios = []



    def set_data(self, date: str):
        self.date = date

    def set_event_name(self, event_name: str):
        self.event_name = event_name

    def set_institute(self, institut_num: int):
        """Определяет институт по его номеру"""
        self.institut_name = DocumentsConfigs.institut_name(institut_num)
        self.director_name = DocumentsConfigs.director_name(institut_num)

    def add_fios(self, fios: list[list[str]]):
        self.fios = fios

    def add_fio(self, fio: list[str]):
        self.fios.append(fio)

    def _add_exemption_content(self):
        # добавление института
        self.exemption_doc.paragraphs[1].add_run(self.institut_name)

        # добавление фио директора
        self.exemption_doc.paragraphs[2].add_run(self.director_name)
        # добавление названия мероприятия
        for i in self.exemption_doc.paragraphs:
            i.text = i.text.replace("event", self.event_name)
        # добавление даты мероприятия
        for j in self.exemption_doc.paragraphs:
            j.text = j.text.replace("date", self.date)
        # Заполнение таблицы
        fio = self.exemption_doc.tables[0]
        for student_fio, group in self.fios:
            num_rows = len(fio.rows)
            cells = fio.add_row().cells
            cells[0].text = str(num_rows) + "."  # порядковый номер
            cells[1].text = student_fio  # ФИО
            cells[2].text = group  # Группа

    def make_exemption(self):
        self.exemption_doc = Document(DocumentsConfigs.sample_exemption())
        style = self.exemption_doc.styles[self._doc_style]
        style.font.name = self._font_style
        style.font.size = Pt(self._font_size)
        self._add_exemption_content()
        self.exemption_doc.save(
            f"Освобождение {self.event_name} {self.date}.docx")

    def _add_gatitude_content(self):
        run = self.thanks_doc.paragraphs[3].add_run(self.institut_name)
        run.bold = True
        fio = self.thanks_doc.tables[2]

        for student_fio, group in self.fios:
            num_rows = len(self.fios)
            cells = fio.add_row().cells
            cells[0].text = str(num_rows) + '.'  # порядковый номер
            cells[1].text = student_fio  # фио студента
            cells[2].text = group  # группа студента

    def make_thanks(self):
        self.thanks_doc = Document(DocumentsConfigs.sample_thanks())
        style = self.thanks_doc.styles[self._doc_style]
        style.font.name = self._font_style
        style.font.size = Pt(self._font_size)
        self._add_gatitude_content()
        self.thanks_doc.save(f"Благодарность  {self.event_name} {self.date}.docx")


# def exemption():
#     doc = Document("SampleExemption.docx")  # макет документа в папке с кодом
#
#     # стили шрифта
#     style = doc.styles['Normal']
#     style.font.name = 'Times New Roman'
#     style.font.size = Pt(14)
#
#     # добавление института в строку ворд
#     institute = doc.paragraphs[1]
#     institute.add_run(input("Введите институт: "))
#
#     # Ввод И. О. Фамилии
#     iof = doc.paragraphs[2]
#     iof.add_run(input("Введите И.О.Фамилию: "))
#
#     # Ввод названия мероприятия
#     event_name = input("Название мероприятия: ")
#
#     for i in doc.paragraphs:
#         i.text = i.text.replace("event", event_name)
#
#     # Ввод даты мероприятия
#     date_event = input("Дата мероприятия: ")
#
#     for j in doc.paragraphs:
#         j.text = j.text.replace("date", date_event)
#
#     # Заполнение таблицы
#     fio = doc.tables[0]
#     while True:
#         num_rows = len(fio.rows)
#         cells = fio.add_row().cells
#         cells[0].text = str(num_rows) + (".")  # порядковый номер
#         cells[1].text = input("Введите ФИО: ")  # ФИО
#         cells[2].text = input("Введите Группу: ")  # Группа
#         next_step = input("Для продолжения нажмите Enter, Чтобы закончить напишите 'e', : ")
#         if next_step.lower() == "e" or next_step.lower() == "е":  # e на русском и на англ
#             break
#
#     # Сохранение документа
#     doc.save(
#         f"ReadyExemption\\Освобождение {event_name} {date_event}.docx")  # необходимо создать папку "ReadyExemption" в корне проекта


new_doc = Documents()
new_doc.set_event_name('Создание чертового автокомпилятора доков')
new_doc.set_data('24.02.2024')
new_doc.set_institute(6)
new_doc.add_fio(['Создатели гладос', 'крови на рукаве'])
new_doc.make_thanks()
